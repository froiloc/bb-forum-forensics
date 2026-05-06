# =============================================================================
# forensic_api/export.py
# IT-Forensisches Ermittlungswerkzeug — Baustelle 6: Berichte & Exports
# =============================================================================
# Zweck:
#   Endpunkt GET /_forensic/export?format=<html|docx|sqlite>
#
#   Erzeugt einen serverseitigen Download in einem der drei Formate:
#
#   Format A (html):
#     Selbstenthaltendes HTML-Dokument. Alle Paragraphen mit
#     status IN ('active', 'approved'), Beweisanker als Fussnoten,
#     Kopf- und Fusszeile. Keine externe Bibliothek erforderlich.
#     Beleg: Bauplan B6 v0.3 §7.2 Format A
#
#   Format B (docx):
#     Word-Dokument (python-docx). Gleicher Inhalt wie HTML,
#     aber mit Word-nativer Formatierung: Deckblatt, Ueberschriften,
#     Absatzstile, Seitennummerierung.
#     Beleg: Bauplan B6 v0.3 §7.2 Format B
#
#   Format C (sqlite):
#     Eigenstaendige SQLite3-Datenbank fallakte_<uid>_<datum>.db.
#     Enthaelt meta, profile_summary, known_aliases, report_paragraphs,
#     evidence_annotations, network_summary, activity_stats,
#     timeline_summary und README.
#     Beleg: Bauplan B6 v0.3 §7.2 Format C
#
# Zugriffssteuerung:
#   Nur eigener Bericht (context.user_id) abrufbar.
#
# Beleg: Bauplan B6 v0.3 §7, Build 097
# Version: v0.6.097 · Build: 097 · 2026-05-05
# =============================================================================

from __future__ import annotations

import io
import json
import sqlite3
import time
from datetime import datetime
from typing import TYPE_CHECKING

from core.logger import get_logger

if TYPE_CHECKING:
    from server.http_server import ForensicRequestHandler
    from db.connection_manager import DatabaseBundle
    from core.config_loader import ConfigLoader
    from core.mode_resolver import ResolvedContext

logger = get_logger(__name__)

# Nur active und approved Paragraphen werden exportiert
EXPORT_STATUSES = ("active", "approved")

# Klassifizierungs-Kennzeichnung
CLASSIFICATION = "VERTRAULICH — IT-FORENSISCHES ERMITTLUNGSWERKZEUG NRW"


def _err(msg: str, code: int = 400) -> tuple[int, bytes]:
    body = json.dumps({"error": msg}, ensure_ascii=False).encode("utf-8")
    return code, body


class ExportEndpoint:
    """
    GET /_forensic/export?format=html|docx|sqlite
    Serverseitiger Datei-Export.
    Beleg: Bauplan B6 v0.3 §7.2, Build 097
    """

    def __init__(
        self,
        bundle: "DatabaseBundle",
        context: "ResolvedContext",
        config: "ConfigLoader",
    ) -> None:
        self._bundle  = bundle
        self._context = context
        self._config  = config

    # ------------------------------------------------------------------
    # Dispatch
    # ------------------------------------------------------------------

    def handle_get(
        self,
        handler: "ForensicRequestHandler",
        params: dict,
    ) -> None:
        fmt = params.get("format", [None])[0]
        if fmt not in ("html", "docx", "sqlite"):
            handler.send_response_body(
                400,
                json.dumps({"error": "format muss html, docx oder sqlite sein"},
                           ensure_ascii=False).encode("utf-8"),
                content_type="application/json; charset=utf-8",
            )
            return

        if fmt == "html":
            self._export_html(handler)
        elif fmt == "docx":
            self._export_docx(handler)
        elif fmt == "sqlite":
            self._export_sqlite(handler)

    # ------------------------------------------------------------------
    # Gemeinsame Datensammlung
    # ------------------------------------------------------------------

    def _collect_paragraphs(self) -> list:
        """
        Laedt alle exportierbaren Paragraphen (active + approved)
        in sort_index-Reihenfolge.
        """
        edb = self._bundle.evidence
        reports = edb.get_reports()
        if not reports:
            return []
        # Erster aktiver Bericht
        report = next(
            (r for r in reports if r.status in ("draft", "submitted")),
            reports[0]
        )
        paras = edb.get_paragraphs(report.id)
        return [p for p in paras if p.status in EXPORT_STATUSES]

    def _collect_anchors(self, block_id: str) -> list:
        """Laedt alle Beweisanker fuer einen Paragraph."""
        return self._bundle.evidence.get_anchors_for_paragraph(block_id)

    def _now_str(self) -> str:
        return datetime.now().strftime("%d.%m.%Y %H:%M")

    def _date_str(self) -> str:
        return datetime.now().strftime("%Y%m%d")

    def _username(self) -> str:
        return self._context.username or f"uid_{self._context.user_id}"

    def _uid(self) -> int:
        return self._context.user_id

    # ------------------------------------------------------------------
    # Format A: HTML
    # Beleg: Bauplan B6 v0.3 §7.2 Format A
    # ------------------------------------------------------------------

    def _export_html(self, handler: "ForensicRequestHandler") -> None:
        paras   = self._collect_paragraphs()
        now_str = self._now_str()
        uname   = self._username()
        uid     = self._uid()

        paragraphs_html = []
        for i, p in enumerate(paras, 1):
            anchors = self._collect_anchors(p.block_id)
            anchor_notes = ""
            if anchors:
                notes = "".join(
                    f"<li>[{a.id}] {self._html_esc(a.anchor_text)}</li>"
                    for a in anchors
                )
                anchor_notes = f'<ol class="anchors">{notes}</ol>'

            paragraphs_html.append(
                f'<div class="paragraph">'
                f'<p class="para-content">{self._html_esc(p.content)}</p>'
                f'{anchor_notes}'
                f'</div>'
            )

        body_html = "\n".join(paragraphs_html) if paragraphs_html else "<p><em>Keine Absaetze vorhanden.</em></p>"

        html = f"""<!DOCTYPE html>
<html lang="de">
<head>
  <meta charset="utf-8">
  <title>Bericht {self._html_esc(uname)} (ID: {uid})</title>
  <style>
    body {{ font-family: "Times New Roman", Times, serif; font-size: 11pt;
           max-width: 800px; margin: 40px auto; color: #000; }}
    .header {{ font-size: 9pt; color: #555; border-bottom: 1pt solid #ccc;
               padding-bottom: 6pt; margin-bottom: 24pt; font-family: Arial, sans-serif; }}
    .paragraph {{ margin-bottom: 18pt; page-break-inside: avoid; }}
    .para-content {{ line-height: 1.6; white-space: pre-wrap; word-break: break-word; }}
    .anchors {{ font-size: 8.5pt; color: #555; margin-top: 4pt; }}
    .footer {{ font-size: 8pt; color: #888; border-top: 1pt solid #ddd;
               padding-top: 4pt; margin-top: 48pt; text-align: center;
               font-family: Arial, sans-serif; font-weight: bold; }}
  </style>
</head>
<body>
  <div class="header">
    Beschuldigter: {self._html_esc(uname)} &middot; ID: {uid}<br>
    Erstellt am: {self._html_esc(now_str)} &middot; {len(paras)} Abs&auml;tze
  </div>
  {body_html}
  <div class="footer">{self._html_esc(CLASSIFICATION)}</div>
</body>
</html>"""

        filename = f"bericht_{uid}_{self._date_str()}.html"
        body = html.encode("utf-8")
        handler.send_response_body(
            200, body,
            content_type="text/html; charset=utf-8",
            extra_headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
                "Content-Length": str(len(body)),
            },
        )
        logger.info(
            "Export HTML: uid=%d, %d Paragraphen, %d Bytes",
            uid, len(paras), len(body),
        )

    # ------------------------------------------------------------------
    # Format B: DOCX
    # Beleg: Bauplan B6 v0.3 §7.2 Format B
    # ------------------------------------------------------------------

    def _export_docx(self, handler: "ForensicRequestHandler") -> None:
        try:
            from docx import Document                       # type: ignore
            from docx.shared import Pt, Cm                 # type: ignore
            from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
        except ImportError:
            handler.send_response_body(
                503,
                json.dumps(
                    {"error": "python-docx nicht installiert. "
                              "Bitte 'pip install python-docx' ausfuehren."},
                    ensure_ascii=False,
                ).encode("utf-8"),
                content_type="application/json; charset=utf-8",
            )
            return

        paras   = self._collect_paragraphs()
        now_str = self._now_str()
        uname   = self._username()
        uid     = self._uid()

        doc = Document()

        # Seitenraender
        for section in doc.sections:
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin   = Cm(3.0)
            section.right_margin  = Cm(2.5)

        # Deckblatt
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(f"Ermittlungsbericht")
        run.bold = True
        run.font.size = Pt(16)

        doc.add_paragraph()
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run(f"Beschuldigter: {uname} (ID: {uid})").font.size = Pt(12)

        doc.add_paragraph()
        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_p.add_run(f"Erstellt: {now_str}").font.size = Pt(10)

        doc.add_page_break()

        # Absaetze
        for i, p in enumerate(paras, 1):
            heading = doc.add_paragraph(style="Heading 2")
            heading.add_run(f"Absatz {i}")

            content_para = doc.add_paragraph()
            content_para.add_run(p.content)

            # Beweisanker als Fussnotenblock
            anchors = self._collect_anchors(p.block_id)
            if anchors:
                fn_para = doc.add_paragraph(style="Normal")
                fn_run  = fn_para.add_run("Beweisanker: ")
                fn_run.bold = True
                fn_run.font.size = Pt(9)
                for a in anchors:
                    fn_para.add_run(f"[{a.id}] {a.anchor_text}  ").font.size = Pt(9)

        # Klassifizierungs-Fusszeile
        footer_section = doc.sections[0]
        footer = footer_section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.add_run(CLASSIFICATION).font.size = Pt(8)

        buf = io.BytesIO()
        doc.save(buf)
        body = buf.getvalue()

        filename = f"bericht_{uid}_{self._date_str()}.docx"
        handler.send_response_body(
            200, body,
            content_type="application/vnd.openxmlformats-officedocument"
                          ".wordprocessingml.document",
            extra_headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
                "Content-Length": str(len(body)),
            },
        )
        logger.info(
            "Export DOCX: uid=%d, %d Paragraphen, %d Bytes",
            uid, len(paras), len(body),
        )

    # ------------------------------------------------------------------
    # Format C: SQLite3
    # Beleg: Bauplan B6 v0.3 §7.2 Format C
    # ------------------------------------------------------------------

    def _export_sqlite(self, handler: "ForensicRequestHandler") -> None:
        uid       = self._uid()
        uname     = self._username()
        now_ts    = int(time.time())
        now_str   = self._now_str()
        paras     = self._collect_paragraphs()
        anns      = self._bundle.evidence.get_all_annotations()
        # system_username ist kein Standardfeld von ResolvedContext --
        # Fallback auf username des Beschuldigten (fuer Tests und Support-Modus)
        _inv_raw = getattr(self._context, 'system_username', None)
        inv      = str(_inv_raw) if _inv_raw and not hasattr(_inv_raw, '_mock_name') else uname

        # In-Memory-SQLite3-Export-Datenbank aufbauen
        con = sqlite3.connect(":memory:")

        # ------ meta ------
        con.execute("""
            CREATE TABLE meta (
                key   TEXT PRIMARY KEY,
                value TEXT NOT NULL
            )""")
        for k, v in [
            ("uid",                    str(uid)),
            ("username",               uname),
            ("export_date",            now_str),
            ("exporting_investigator", inv),
            ("tool_version",           "aiw_webserver v0.6.097"),
            ("classification",         CLASSIFICATION),
        ]:
            con.execute("INSERT INTO meta VALUES (?, ?)", (k, v))

        # ------ profile_summary ------
        con.execute("""
            CREATE TABLE profile_summary (
                -- Wichtigste Profildaten des Beschuldigten.
                -- Kein Passwort, kein session_id.
                key   TEXT PRIMARY KEY,
                value TEXT
            )""")
        try:
            fcon = self._bundle.connection
            row  = fcon.execute(
                "SELECT username, email, registered, group_id "
                "FROM fdb.uid_profile WHERE id = ?", (uid,)
            ).fetchone()
            if row:
                for k, v in [
                    ("username",    str(row[0] or "")),
                    ("email",       str(row[1] or "")),
                    ("registered",  str(row[2] or "")),
                    ("group_id",    str(row[3] or "")),
                ]:
                    con.execute("INSERT INTO profile_summary VALUES (?, ?)", (k, v))
        except Exception as exc:
            logger.debug("SQLite-Export profile_summary: %s", exc)
            con.execute("INSERT INTO profile_summary VALUES (?, ?)",
                        ("note", "Profildaten nicht verfuegbar"))

        # ------ known_aliases ------
        con.execute("""
            CREATE TABLE known_aliases (
                -- Alle bekannten Aliasnamen des Beschuldigten aus uid_aliases.
                historical_username TEXT NOT NULL,
                first_seen_ts       INTEGER,
                last_seen_ts        INTEGER,
                source              TEXT
            )""")
        try:
            rows = self._bundle.connection.execute(
                "SELECT historical_username, first_seen_ts, last_seen_ts, source "
                "FROM fdb.uid_aliases"
            ).fetchall()
            for row in rows:
                con.execute("INSERT INTO known_aliases VALUES (?, ?, ?, ?)", tuple(row))
        except Exception as exc:
            logger.debug("SQLite-Export known_aliases: %s", exc)

        # ------ report_paragraphs ------
        con.execute("""
            CREATE TABLE report_paragraphs (
                -- Berichtsabsaetze mit Status 'active' oder 'approved'.
                -- Entwuerfe, ausgelassene und ersetzte Absaetze sind nicht enthalten.
                block_id   TEXT    NOT NULL PRIMARY KEY,
                content    TEXT    NOT NULL,
                status     TEXT    NOT NULL,
                author     TEXT    NOT NULL,
                created_at INTEGER NOT NULL,
                updated_at INTEGER NOT NULL
            )""")
        for p in paras:
            con.execute(
                "INSERT INTO report_paragraphs VALUES (?, ?, ?, ?, ?, ?)",
                (p.block_id, p.content, p.status, p.author,
                 p.created_at, p.updated_at)
            )

        # ------ evidence_annotations ------
        con.execute("""
            CREATE TABLE evidence_annotations (
                -- Alle Annotationen des Ermittlers (abgeschlossene Markierungen).
                id          INTEGER NOT NULL PRIMARY KEY,
                page_url    TEXT    NOT NULL,
                category    TEXT    NOT NULL,
                note_text   TEXT,
                orig_text   TEXT,
                created_by  TEXT,
                created_at  INTEGER
            )""")
        for a in anns:
            orig = None
            if a.selection_json:
                try:
                    sel  = json.loads(a.selection_json)
                    orig = sel.get("text") if isinstance(sel, dict) else None
                except Exception:
                    pass
            con.execute(
                "INSERT INTO evidence_annotations VALUES (?, ?, ?, ?, ?, ?, ?)",
                (a.id, a.page_url, a.category, a.text or None,
                 orig, a.created_by or None, a.ts)
            )

        # ------ network_summary ------
        con.execute("""
            CREATE TABLE network_summary (
                -- Kommunikationspartner des Beschuldigten aus uid_pn_network.
                partner_user_id   INTEGER,
                partner_username  TEXT,
                message_count     INTEGER
            )""")
        try:
            rows = self._bundle.connection.execute(
                "SELECT partner_user_id, partner_username, message_count "
                "FROM fdb.uid_pn_network"
            ).fetchall()
            for row in rows:
                con.execute("INSERT INTO network_summary VALUES (?, ?, ?)", tuple(row))
        except Exception as exc:
            logger.debug("SQLite-Export network_summary: %s", exc)

        # ------ activity_stats ------
        con.execute("""
            CREATE TABLE activity_stats (
                -- Aktivitaetskennzahlen des Beschuldigten aus uid_stats.
                stat_key       TEXT NOT NULL PRIMARY KEY,
                val_computed   TEXT,
                display_label  TEXT
            )""")
        try:
            rows = self._bundle.connection.execute(
                "SELECT stat_key, val_computed, display_label FROM fdb.uid_stats"
            ).fetchall()
            for row in rows:
                con.execute("INSERT INTO activity_stats VALUES (?, ?, ?)", tuple(row))
        except Exception as exc:
            logger.debug("SQLite-Export activity_stats: %s", exc)

        # ------ timeline_summary ------
        con.execute("""
            CREATE TABLE timeline_summary (
                -- Monatliche Aktivitaetsuebersicht aus uid_timeline_agg.
                month_key      TEXT    NOT NULL PRIMARY KEY,
                post_count     INTEGER,
                topic_count    INTEGER
            )""")
        try:
            rows = self._bundle.connection.execute(
                "SELECT month_key, post_count, topic_count FROM fdb.uid_timeline_agg"
            ).fetchall()
            for row in rows:
                con.execute("INSERT INTO timeline_summary VALUES (?, ?, ?)", tuple(row))
        except Exception as exc:
            logger.debug("SQLite-Export timeline_summary: %s", exc)

        # ------ README ------
        con.execute("""
            CREATE TABLE README (
                -- Menschenlesbare Beschreibung der Datenbank und ihrer Tabellen.
                key   TEXT PRIMARY KEY,
                value TEXT NOT NULL
            )""")
        readme_entries = [
            ("beschreibung",
             "Fallakte des IT-forensischen Ermittlungswerkzeugs NRW. "
             "Diese Datenbank enthaelt alle Ermittlungsergebnisse zu einem Beschuldigten "
             "und ist selbsterlaeuternd konzipiert."),
            ("tabelle_meta",
             "Metadaten: uid, Benutzername, Exportzeitpunkt, exportierender Ermittler, "
             "Werkzeugversion, Klassifizierung."),
            ("tabelle_profile_summary",
             "Wichtigste Profildaten des Beschuldigten (kein Passwort, keine Session-ID)."),
            ("tabelle_known_aliases",
             "Alle bekannten historischen Benutzernamen des Beschuldigten."),
            ("tabelle_report_paragraphs",
             "Freigegebene und aktive Berichtsabsaetze. Status 'draft', 'omitted' und "
             "'superseded' sind nicht enthalten."),
            ("tabelle_evidence_annotations",
             "Annotationen des Ermittlers: Markierungen, Kategorisierungen, Notizen."),
            ("tabelle_network_summary",
             "Kommunikationspartner des Beschuldigten im Forum (Privatnachrichten)."),
            ("tabelle_activity_stats",
             "Aktivitaetskennzahlen: Beitragsanzahl, PN-Partner, Dankesagungen etc."),
            ("tabelle_timeline_summary",
             "Monatliche Aktivitaetsuebersicht des Beschuldigten."),
            ("klassifizierung", CLASSIFICATION),
        ]
        for k, v in readme_entries:
            con.execute("INSERT INTO README VALUES (?, ?)", (k, v))

        con.commit()

        # SQLite-Datenbank als Bytes serialisieren
        # Methode: serialize() ist ab Python 3.12 verfuegbar.
        # Fallback fuer aeltere Python-Versionen: tempfile.
        try:
            body = con.serialize()
        except AttributeError:
            import tempfile, os
            with tempfile.NamedTemporaryFile(suffix=".db", delete=False) as tmp:
                tmp_path = tmp.name
            try:
                backup_con = sqlite3.connect(tmp_path)
                con.backup(backup_con)
                backup_con.close()
                with open(tmp_path, "rb") as f:
                    body = f.read()
            finally:
                os.unlink(tmp_path)
        finally:
            con.close()

        filename = f"fallakte_{uid}_{self._date_str()}.db"
        handler.send_response_body(
            200, body,
            content_type="application/octet-stream",
            extra_headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
                "Content-Length": str(len(body)),
            },
        )
        logger.info(
            "Export SQLite: uid=%d, %d Paragraphen, %d Annotationen, %d Bytes",
            uid, len(paras), len(anns), len(body),
        )

    # ------------------------------------------------------------------
    # Hilfsfunktionen
    # ------------------------------------------------------------------

    @staticmethod
    def _html_esc(s: str) -> str:
        if not s:
            return ""
        return (str(s)
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
                .replace('"', "&quot;"))
