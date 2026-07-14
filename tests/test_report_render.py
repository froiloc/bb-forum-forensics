# =============================================================================
# tests/test_report_render.py
# IT-Forensisches Ermittlungswerkzeug — Baustelle 6/7: Berichts-Ausgabe
# =============================================================================
# Testsuite fuer das serverunabhaengige Paket report_render/ (Build 399).
#
# WICHTIG (Bauplan Build 397 §5): Die Tests laufen gegen eine ECHTE EvidenceDb
# auf einer temporaeren Datei — KEIN MagicMock an der Stelle, an der das
# Interface geprueft wird. Genau dort war der Fehler B1/B2 entstanden
# ("gruen aber tot").
#
# Abdeckung:
#   Resolver-Paritaet:   PR01-PR06  (a/m/o, XSS, unbekannt, Newline)
#   Resolver Text-Modus: TX01-TX04  (Tags/Entities, roh vs. escaped, <br>, 399-Paritaet)
#   ReportSource:        RS01-RS06  (Auswahl, Warnungen R2, Bild-Verweis, NoReport)
#   Bild-Anreicherung:   IE01       (BLOB-freie Referenzfelder, Build 402)
#   HtmlRenderer:        HR01-HR07  (Statuskopf R1, Bloecke, R2-Hinweise, R3)
#   DocxRenderer:        DX01-DX03  (Status/Bloecke, R3, XSS-als-Text) [Build 402]
#   SqliteRenderer:      SQ01-SQ04  (Tabellen/Zahlen, R2-Warnungen, Bild-Verweis, R3) [Build 402]
#
# Version: v0.7.402 · Build: 402 · 2026-07-14
# =============================================================================

import json
import sqlite3
import sys
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from db.evidence_db import EvidenceDb
from report_render.placeholder_resolver import PlaceholderResolver
from report_render.report_source import ReportSource, NoReportError
from report_render.html_renderer import HtmlRenderer
from report_render.report_document import (
    WARN_UNRESOLVED_PLACEHOLDER, WARN_UNKNOWN_PLACEHOLDER,
    WARN_UNORDERED_BLOCK, WARN_UNKNOWN_BLOCK_TYPE, WARN_MISSING_IMAGE,
)


# -----------------------------------------------------------------------------
# Fixture: echte EvidenceDb mit einem reichhaltigen Bericht.
# -----------------------------------------------------------------------------
def _make_edb():
    con = sqlite3.connect(":memory:", check_same_thread=False)
    return con, EvidenceDb(con)


def _seed(con):
    """Legt zwei Berichte an; der hoehere sequence_nr (id=2) ist zu waehlen."""
    # report 1: niedrigere sequence_nr -> darf NICHT gewaehlt werden
    con.execute("INSERT INTO reports (id, report_type, sequence_nr, title, created_by, created_at, status) "
                "VALUES (1,'interim',1,'Alt','inv',1000,'draft')")
    # report 2: hoechste sequence_nr -> wird gewaehlt (mc §4.1)
    con.execute("INSERT INTO reports (id, report_type, sequence_nr, title, created_by, created_at, status) "
                "VALUES (2,'final',3,'Haupt','inv',2000,'submitted')")

    def blk(bid, btype, data, values=None):
        con.execute(
            "INSERT INTO report_blocks (block_id, report_id, author, created_at, updated_at, "
            "block_type, block_data, placeholder_values_json, module_id) "
            "VALUES (?,?,?,?,?,?,?,?,?)",
            (bid, 2, "inv", 2000, 2000, btype, json.dumps(data),
             json.dumps(values) if values is not None else None, None),
        )

    def order(bid, idx):
        con.execute("INSERT INTO report_block_order (block_id, sort_index, last_modified_by, last_modified_at) "
                    "VALUES (?,?,?,?)", (bid, idx, "inv", 2000))

    # Paragraph mit a/m/o-Platzhaltern; m-Wert enthaelt XSS-Nutzlast.
    blk("b_para", "paragraph",
        {"text": "Name {{a:user.name}}, Zeit {{m:tatzeit|}}, Ort {{o:ort|}}. Ende."},
        {"tatzeit": "<script>alert(1)</script>", "ort": ""})
    blk("b_head", "header", {"text": "Kapitel", "level": 2})
    blk("b_list", "list", {"style": "unordered", "items": ["Eins", "Zwei"]})
    blk("b_table", "table", {"withHeadings": True, "content": [["A", "B"], ["1", "2"]]})
    blk("b_quote", "quote", {"text": "Zitat", "caption": "Quelle"})
    blk("b_img", "image", {"url": "/img/avatars/1.jpg", "caption": "Profilbild"})
    blk("b_delim", "delimiter", {})
    blk("b_marker", "marker", {"text": "Hervorgehoben"})
    blk("b_evidence", "evidence", {"evidence_ids": [10, 11], "text": "Beweislage"})
    blk("b_unknown", "audio", {"text": "kann nicht dargestellt werden"})  # R3
    blk("b_unordered", "paragraph", {"text": "Ohne Ordnung"})            # R2: kein order-Eintrag

    for i, bid in enumerate(
        ["b_para", "b_head", "b_list", "b_table", "b_quote", "b_img",
         "b_delim", "b_marker", "b_evidence", "b_unknown"]):
        order(bid, i)
    # b_unordered ABSICHTLICH ohne order-Eintrag.

    # Ein Beweisanker am Paragraph.
    con.execute("INSERT INTO report_anchors (id, block_id, annotation_id, anchor_text, created_at) "
                "VALUES (1,'b_para',99,'Beleg: Beitrag #5',2000)")
    con.commit()


# -----------------------------------------------------------------------------
# Resolver-Paritaet (Golden Cases aus placeholder_chips.js abgeleitet)
# -----------------------------------------------------------------------------
class TestResolverParity(unittest.TestCase):

    def setUp(self):
        # resolve_auto liefert einen Wert nur fuer 'known', sonst None/"".
        def auto(name):
            return {"known": "Wert", "leer": ""}.get(name, None)
        self.r = PlaceholderResolver(resolve_auto=auto)

    def test_PR01_auto_resolved(self):
        frag, warns = self.r.resolve("X {{a:known}} Y")
        self.assertEqual(frag, "X Wert Y")
        self.assertEqual(warns, [])

    def test_PR02_auto_unresolvable_uses_default_and_warns(self):
        frag, warns = self.r.resolve("{{a:fehlt|Standard}}")
        self.assertEqual(frag, "Standard")
        self.assertEqual(len(warns), 1)
        self.assertEqual(warns[0].kind, WARN_UNRESOLVED_PLACEHOLDER)

    def test_PR03_auto_empty_shows_name_and_warns(self):
        # resolved == "" -> Client zeigt name; R2 -> Warnung.
        frag, warns = self.r.resolve("{{a:leer}}")
        self.assertEqual(frag, "leer")
        self.assertEqual(len(warns), 1)

    def test_PR04_mandatory_filled_is_escaped(self):
        frag, warns = self.r.resolve("{{m:t|}}", {"t": "<b>&\"x"})
        self.assertEqual(frag, "&lt;b&gt;&amp;&quot;x")
        self.assertEqual(warns, [])

    def test_PR05_mandatory_empty_warns_optional_empty_silent(self):
        frag_m, warns_m = self.r.resolve("{{m:pflicht|}}", {})
        self.assertIn("*", frag_m)
        self.assertEqual(len(warns_m), 1)
        self.assertEqual(warns_m[0].kind, WARN_UNRESOLVED_PLACEHOLDER)
        frag_o, warns_o = self.r.resolve("{{o:opt|}}", {})
        self.assertEqual(warns_o, [])            # optional leer -> keine Warnung

    def test_PR06_unknown_braces_and_newline(self):
        frag, warns = self.r.resolve("A\n{{x:foo}} B")
        self.assertIn("<br>", frag)               # \n -> <br>
        self.assertIn("{{x:foo}}", frag)          # unveraendert (Paritaet)
        self.assertTrue(any(w.kind == WARN_UNKNOWN_PLACEHOLDER for w in warns))


# -----------------------------------------------------------------------------
# ReportSource
# -----------------------------------------------------------------------------
class TestReportSource(unittest.TestCase):

    def setUp(self):
        self.con, self.edb = _make_edb()
        _seed(self.con)
        self.src = ReportSource(
            evidence=self.edb, templates=None, assets=None, forensic_con=None,
            uid=42, username="TestNutzer", generated_at=1_700_000_000,
        )

    def test_RS01_selects_highest_sequence_nr(self):
        doc = self.src.build()
        self.assertEqual(doc.report_id, 2)
        self.assertEqual(doc.sequence_nr, 3)
        self.assertEqual(doc.status, "submitted")

    def test_RS02_explicit_report_id(self):
        doc = self.src.build(report_id=1)
        self.assertEqual(doc.report_id, 1)

    def test_RS03_all_blocks_present_including_unordered(self):
        doc = self.src.build()
        self.assertEqual(len(doc.blocks), 11)     # 10 geordnet + 1 ungeordnet

    def test_RS04_warnings_cover_all_kinds(self):
        doc = self.src.build()
        kinds = {w.kind for w in doc.warnings}
        self.assertIn(WARN_UNORDERED_BLOCK, kinds)      # b_unordered
        self.assertIn(WARN_UNKNOWN_BLOCK_TYPE, kinds)   # b_unknown (audio)
        self.assertIn(WARN_MISSING_IMAGE, kinds)        # assets=None
        self.assertIn(WARN_UNRESOLVED_PLACEHOLDER, kinds)  # {{a:user.name}}

    def test_RS05_image_is_reference_not_bytes(self):
        doc = self.src.build()
        img = next(b for b in doc.blocks if b.block_id == "b_img")
        self.assertEqual(img.data["_image_url"], "/img/avatars/1.jpg")
        self.assertFalse(img.data["_image_available"])

    def test_RS06_no_report_raises(self):
        con2, edb2 = _make_edb()      # leere DB
        src2 = ReportSource(edb2, None, None, None, 1, "x", 1)
        with self.assertRaises(NoReportError):
            src2.build()


# -----------------------------------------------------------------------------
# HtmlRenderer
# -----------------------------------------------------------------------------
class TestHtmlRenderer(unittest.TestCase):

    def setUp(self):
        self.con, self.edb = _make_edb()
        _seed(self.con)
        self.src = ReportSource(self.edb, None, None, None, 42, "TestNutzer", 1_700_000_000)

    def _html(self):
        return HtmlRenderer().render(self.src.build()).decode("utf-8")

    def test_HR01_valid_utf8_html(self):
        html = self._html()
        self.assertTrue(html.startswith("<!DOCTYPE html>"))
        self.assertIn('<meta charset="utf-8">', html)

    def test_HR02_status_banner_submitted(self):
        self.assertIn("ZUR ABNAHME VORGELEGT", self._html())

    def test_HR03_xss_payload_escaped(self):
        html = self._html()
        self.assertNotIn("<script>alert(1)</script>", html)
        self.assertIn("&lt;script&gt;alert(1)&lt;/script&gt;", html)

    def test_HR04_unknown_block_visible(self):
        html = self._html()
        self.assertIn("Unbekannter Blocktyp 'audio'", html)

    def test_HR05_image_reference_not_embedded(self):
        html = self._html()
        self.assertIn("nicht eingebettet", html)
        self.assertNotIn("<img", html)            # KEINE Bild-Einbettung (§4.2)

    def test_HR06_hints_section_lists_warnings(self):
        html = self._html()
        self.assertIn("Hinweise zur Erzeugung", html)
        self.assertIn("Warnungen", html)

    def test_HR07_anchor_and_table_rendered(self):
        html = self._html()
        self.assertIn("Beleg: Beitrag #5", html)   # Anker
        self.assertIn("<table class=\"report-table\">", html)
        self.assertIn("<th>A</th>", html)          # withHeadings -> erste Zeile th


# -----------------------------------------------------------------------------
# Resolver Text-Modus (Build 402)
# -----------------------------------------------------------------------------
class TestResolverTextMode(unittest.TestCase):

    def setUp(self):
        self.r = PlaceholderResolver(resolve_auto=lambda n: {"k": "V"}.get(n))

    def test_TX01_text_mode_strips_tags_unescapes(self):
        # Editor.js-HTML im Textsegment -> Tags weg, Entities aufgeloest.
        frag, _ = self.r.resolve("<b>Hallo</b> &amp; Welt", mode="text")
        self.assertEqual(frag, "Hallo & Welt")

    def test_TX02_text_mode_value_raw_html_mode_escaped(self):
        h, t, _ = self.r.resolve_both("{{m:x|}}", {"x": "<i>&"})
        self.assertEqual(t, "<i>&")                    # roh im Text
        self.assertEqual(h, "&lt;i&gt;&amp;")          # escaped im HTML

    def test_TX03_br_becomes_newline_in_text(self):
        frag, _ = self.r.resolve("A<br>B", mode="text")
        self.assertEqual(frag, "A\nB")

    def test_TX04_html_mode_identical_to_build399(self):
        # Regressionsschutz: HTML-Serialisierung unveraendert.
        frag, _ = self.r.resolve("X {{a:k}} Y\nZ", mode="html")
        self.assertEqual(frag, "X V Y<br>Z")


# -----------------------------------------------------------------------------
# Bild-Anreicherung (BLOB-frei, Build 402)
# -----------------------------------------------------------------------------
class _FakeAssets:
    """AssetsDb-Ersatz mit get_asset_reference (BLOB-frei)."""
    def __init__(self, known):
        self._known = known
    def get_asset_reference(self, url):
        if url in self._known:
            return {"url_hash": "H" + url, "asset_id": 7, "mime_type": "image/jpeg", "file_size": 123}
        return None


class TestImageEnrichment(unittest.TestCase):

    def test_IE01_reference_fields_present_when_known(self):
        con, edb = _make_edb()
        _seed(con)
        src = ReportSource(edb, None, _FakeAssets({"/img/avatars/1.jpg"}), None,
                           42, "u", 1_700_000_000)
        doc = src.build()
        img = next(b for b in doc.blocks if b.block_id == "b_img")
        self.assertTrue(img.data["_image_available"])
        self.assertEqual(img.data["_image_url_hash"], "H/img/avatars/1.jpg")
        self.assertEqual(img.data["_image_asset_id"], 7)
        self.assertEqual(img.data["_image_size"], 123)
        # keine missing_image-Warnung
        self.assertFalse(any(w.kind == WARN_MISSING_IMAGE for w in doc.warnings))


# -----------------------------------------------------------------------------
# DocxRenderer (Build 402)
# -----------------------------------------------------------------------------
class TestDocxRenderer(unittest.TestCase):

    def setUp(self):
        self.con, self.edb = _make_edb()
        _seed(self.con)
        self.src = ReportSource(self.edb, None, None, None, 42, "TestNutzer", 1_700_000_000)

    def _text(self):
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx nicht installiert")
        import io
        from report_render.docx_renderer import DocxRenderer
        body = DocxRenderer().render(self.src.build())
        self.assertEqual(body[:2], b"PK")            # ZIP-Container
        d = Document(io.BytesIO(body))
        parts = [p.text for p in d.paragraphs]
        for t in d.tables:
            for row in t.rows:
                for c in row.cells:
                    parts.append(c.text)
        return "\n".join(parts)

    def test_DX01_status_and_blocks(self):
        txt = self._text()
        self.assertIn("ZUR ABNAHME VORGELEGT", txt)          # R1
        self.assertIn("Kapitel", txt)                        # header
        self.assertIn("Bildverweis", txt)                    # image ref (§4.2)
        self.assertIn("Hinweise zur Erzeugung", txt)         # R2

    def test_DX02_unknown_block_reported(self):
        self.assertIn("Unbekannter Blocktyp 'audio'", self._text())  # R3

    def test_DX03_xss_is_plain_text_not_markup(self):
        # In DOCX ist der Wert reiner Text; er darf als Zeichenfolge erscheinen.
        self.assertIn("<script>alert(1)</script>", self._text())


# -----------------------------------------------------------------------------
# SqliteRenderer (Build 402)
# -----------------------------------------------------------------------------
class TestSqliteRenderer(unittest.TestCase):

    def setUp(self):
        self.con, self.edb = _make_edb()
        _seed(self.con)
        self.src = ReportSource(self.edb, None, None, None, 42, "TestNutzer", 1_700_000_000)

    def _open(self):
        import os
        import tempfile
        from report_render.sqlite_renderer import SqliteRenderer
        body = SqliteRenderer().render(self.src.build())
        self.assertEqual(body[:16], b"SQLite format 3\x00")
        fd, path = tempfile.mkstemp(suffix=".db")
        os.close(fd)
        with open(path, "wb") as f:
            f.write(body)
        con = sqlite3.connect(path)
        con.row_factory = sqlite3.Row
        self.addCleanup(lambda: (con.close(), os.unlink(path)))
        return con

    def test_SQ01_tables_and_counts(self):
        con = self._open()
        nblocks = con.execute("SELECT COUNT(*) FROM report_blocks").fetchone()[0]
        self.assertEqual(nblocks, 11)
        meta = dict(con.execute("SELECT key, value FROM meta").fetchall())
        self.assertEqual(meta["report_id"], "2")
        self.assertEqual(meta["status"], "submitted")

    def test_SQ02_warnings_persisted(self):
        con = self._open()
        kinds = {r[0] for r in con.execute("SELECT DISTINCT kind FROM report_warnings").fetchall()}
        self.assertIn(WARN_MISSING_IMAGE, kinds)
        self.assertIn(WARN_UNKNOWN_BLOCK_TYPE, kinds)

    def test_SQ03_image_is_reference(self):
        con = self._open()
        row = con.execute(
            "SELECT image_url, image_available FROM report_blocks WHERE block_id='b_img'"
        ).fetchone()
        self.assertEqual(row["image_url"], "/img/avatars/1.jpg")
        self.assertEqual(row["image_available"], 0)   # assets=None -> nicht verfuegbar

    def test_SQ04_unknown_block_flagged(self):
        con = self._open()
        row = con.execute(
            "SELECT is_known_type FROM report_blocks WHERE block_id='b_unknown'"
        ).fetchone()
        self.assertEqual(row["is_known_type"], 0)


if __name__ == "__main__":
    unittest.main()
