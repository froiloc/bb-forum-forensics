# =============================================================================
# report_render/docx_renderer.py
# IT-Forensisches Ermittlungswerkzeug — Baustelle 6/7: Berichts-Ausgabe
# =============================================================================
# Zweck:
#   Rendert ein ReportDocument in ein Word-Dokument (.docx, python-docx).
#   Sieht NUR das ReportDocument, NIE eine Datenbank (Bauplan §2).
#
#   DOCX ist eine REINE-TEXT-Darstellung: es nutzt die vom ReportSource
#   erzeugten *_plain-Felder (Editor.js-HTML entfernt, Entities aufgeloest,
#   Platzhalter aufgeloest). Inline-Formatierung (fett/kursiv) des Editors wird
#   bewusst nicht uebernommen — dasselbe Verhalten wie der Alt-DOCX-Export
#   (v0.6.097), aber jetzt auf dem gemeinsamen Fundament.
#
#   Nicht verhandelbar (Bauplan §3):
#     R1 — Statuszeile sichtbar im Kopf.
#     R2 — Abschnitt "Hinweise zur Erzeugung" mit vollstaendiger Warnliste.
#     R3 — unbekannter Blocktyp sichtbar gemeldet, nicht uebersprungen.
#
#   python-docx wird ERST in render() importiert, damit report_render auch ohne
#   die Bibliothek importierbar bleibt (der Endpunkt meldet ihr Fehlen als 503).
#
# Version: v0.7.402 · Build: 402 · 2026-07-14
# =============================================================================

from __future__ import annotations

import io

from report_render.report_document import ReportDocument, RenderedBlock
# Vorgang 9c41a7e6: die im Werkzeug gewaehlte Zitatvariante.
from report_render.quote_typen import (
    QUOTE_TYP_ANFUEHRUNG, QUOTE_TYP_FELD, QUOTE_TYP_KASTEN, normalisiere,
)

CLASSIFICATION = "VERTRAULICH — IT-FORENSISCHES ERMITTLUNGSWERKZEUG NRW"

# Build 473 (Refactoring "Bericht" -> "Vermerk", mc 2026-07-21): Anzeigelabels
# umbenannt; DB-Schluessel unveraendert (migrationsneutral). Beleg: Auftrag 2026-07-21.
REPORT_TYPE_LABELS = {
    "interim": "Vermerk", "final": "Abschlussbericht", "addendum": "Ergänzungsvermerk",
}
WARNING_LABELS = {
    "unresolved_placeholder": "Nicht auflösbarer Platzhalter (Default eingesetzt)",
    "unknown_placeholder":    "Unbekannter Platzhalter (unveraendert belassen)",
    "unordered_block":        "Block ohne Sortierungseintrag (ans Ende gestellt)",
    "unknown_block_type":     "Unbekannter Blocktyp (Inhalt nicht regulaer dargestellt)",
    "missing_image":          "Bild-Verweis nicht auffindbar",
}


def _absatzrahmen(absatz) -> None:
    """
    Einen duennen Rahmen um einen Absatz legen - die Variante 'Kasten'
    (Vorgang 9c41a7e6).

    WARUM ROHES XML: python-docx bietet fuer Absatzrahmen keine oeffentliche
    Schnittstelle. Ohne diesen Umweg fiele die Variante 'Kasten' im
    Word-Export mit 'senkrechter Strich' zusammen - eine von drei
    Bedienmoeglichkeiten haette dann WEITERHIN keine Wirkung, und genau
    darum geht es in diesem Vorgang.

    DAS RISIKO WIRD BENANNT UND NICHT VERSCHWIEGEN: '_p' ist ein internes
    Attribut von python-docx. Es ist seit 0.8 unveraendert und in 1.2.0
    vorhanden (requirements.txt verlangt >= 1.1.0); bricht es in einer
    kuenftigen Fassung weg, faellt es hier auf - und zwar im Test
    QT13, der den erzeugten Rahmen im XML nachweist, nicht erst beim
    Ermittler.

    Die Masse folgen dem Bildschirm: 'sz' ist in Achtelpunkten (4 = 0,5 pt),
    'space' in Punkten Abstand zum Text, die Farbe ist die des Buendels
    (#d7d7d7).
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = absatz._p.get_or_add_pPr()
    rahmen = OxmlElement("w:pBdr")
    for kante in ("top", "left", "bottom", "right"):
        el = OxmlElement("w:" + kante)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), "D7D7D7")
        rahmen.append(el)
    pPr.append(rahmen)


class DocxRendererUnavailable(RuntimeError):
    """python-docx ist nicht installiert."""


class DocxRenderer:
    """ReportDocument -> bytes (.docx)."""

    def render(self, doc: ReportDocument) -> bytes:
        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError as exc:      # pragma: no cover - Umgebungsabhaengig
            raise DocxRendererUnavailable(str(exc)) from exc

        from datetime import datetime

        d = Document()
        for section in d.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(2.5)

        def _ts(ts):
            try:
                return datetime.fromtimestamp(int(ts)).strftime("%d.%m.%Y %H:%M")
            except (ValueError, OSError, OverflowError):
                return str(ts)

        # --- R1: Statuszeile ---
        st = doc.status
        banner = {
            "draft":     ("ENTWURF — nicht freigegeben. Nicht zur Vorlage bestimmt.", RGBColor(0x7A, 0, 0)),
            "submitted": ("ZUR ABNAHME VORGELEGT — noch nicht freigegeben.", RGBColor(0x7A, 0x52, 0)),
            "approved":  ("FREIGEGEBEN — Stand: " + _ts(doc.generated_at), RGBColor(0x14, 0x53, 0x2D)),
            "final":     ("ABGESCHLOSSEN / AN STA VERSANDT — Stand: " + _ts(doc.generated_at), RGBColor(0x14, 0x53, 0x2D)),
        }.get(st, (f"UNBEKANNTER STATUS: {st} — Behandlung wie Entwurf.", RGBColor(0x7A, 0, 0)))
        p = d.add_paragraph()
        run = p.add_run(banner[0])
        run.bold = True
        run.font.color.rgb = banner[1]

        # --- Titel + Meta ---
        type_label = REPORT_TYPE_LABELS.get(doc.report_type, doc.report_type)
        title = d.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = title.add_run("Ermittlungsbericht")
        tr.bold = True
        tr.font.size = Pt(16)

        meta = d.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(
            f"Beschuldigter: {doc.username} (ID: {doc.uid})\n"
            f"{type_label} Nr. {doc.sequence_nr} — {doc.title}\n"
            f"Erstellt: {_ts(doc.generated_at)} — {len(doc.blocks)} Bloecke — "
            f"{doc.anchor_count} Beweisanker"
        ).font.size = Pt(10)

        d.add_page_break()

        # --- Bloecke ---
        for blk in doc.blocks:
            self._render_block(d, blk, Pt)

        # --- R2: Hinweise zur Erzeugung ---
        d.add_paragraph()
        h = d.add_paragraph()
        h.add_run("Hinweise zur Erzeugung").bold = True
        d.add_paragraph(
            f"Erzeugt: {_ts(doc.generated_at)} — Status: {doc.status} — "
            f"Bloecke: {len(doc.blocks)} — Beweisanker: {doc.anchor_count}"
        )
        if doc.warnings:
            d.add_paragraph(f"Warnungen ({len(doc.warnings)}):")
            for w in doc.warnings:
                label = WARNING_LABELS.get(w.kind, w.kind)
                loc = f" (Block {w.block_id})" if w.block_id else ""
                d.add_paragraph(f"{label}: {w.detail}{loc}", style="List Bullet")
        else:
            d.add_paragraph("Keine Warnungen — alle Platzhalter aufgeloest, "
                            "alle Bloecke sortiert und bekannt.")

        # --- Klassifizierungs-Fusszeile ---
        footer_p = d.sections[0].footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_p.add_run(CLASSIFICATION).font.size = Pt(8)

        buf = io.BytesIO()
        d.save(buf)
        return buf.getvalue()

    # ------------------------------------------------------------------
    def _render_block(self, d, blk: RenderedBlock, Pt) -> None:
        # ORTSNAHER IMPORT wie in render(): python-docx ist eine optionale
        # Abhaengigkeit, und diese Methode wird ausschliesslich aus render()
        # gerufen - also erst, nachdem der dortige Import gelungen ist. Ein
        # Import am Dateikopf machte report_render ohne python-docx
        # unimportierbar und braeche die uebrigen Ausgabeformate mit.
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if not blk.is_known_type:      # R3
            r = d.add_paragraph().add_run(
                f"⚠ Unbekannter Blocktyp '{blk.block_type}' — "
                f"Inhalt nicht regulaer dargestellt (siehe Hinweise)."
            )
            r.bold = True
            self._anchors(d, blk, Pt)
            return

        bt = blk.block_type
        if bt == "paragraph":
            d.add_paragraph(blk.resolved_text_plain)
        elif bt == "header":
            try:
                level = int(blk.data.get("level", 2))
            except (ValueError, TypeError):
                level = 2
            d.add_heading(blk.resolved_text_plain or "", min(max(level, 1), 4))
        elif bt == "list":
            ordered = blk.data.get("style") == "ordered"
            style = "List Number" if ordered else "List Bullet"
            for it in blk.data.get("_resolved_items_plain", []):
                d.add_paragraph(it, style=style)
        elif bt == "table":
            rows = blk.data.get("_resolved_rows_plain", [])
            if rows:
                ncols = max(len(r) for r in rows)
                t = d.add_table(rows=0, cols=ncols)
                try:
                    t.style = "Table Grid"
                except KeyError:       # pragma: no cover
                    pass
                for row in rows:
                    cells = t.add_row().cells
                    for i in range(ncols):
                        cells[i].text = row[i] if i < len(row) else ""
        elif bt == "quote":
            # VORGANG 9c41a7e6: die gewaehlte Zitatvariante bekommt Wirkung.
            # Grundstil bleibt 'Intense Quote' - das ist die Darstellung, die
            # der Bericht bis Build 718 fuer JEDES Zitat hatte. Wer die
            # Variante 'senkrechter Strich' gewaehlt hat, sieht danach also
            # unveraendert dasselbe.
            typ = normalisiere(blk.data.get(QUOTE_TYP_FELD))
            zitat = d.add_paragraph(blk.resolved_text_plain,
                                    style="Intense Quote")
            if typ == QUOTE_TYP_ANFUEHRUNG:
                zitat.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif typ == QUOTE_TYP_KASTEN:
                _absatzrahmen(zitat)
            cap = blk.data.get("_resolved_caption_plain", "")
            if cap:
                # BUILD 719 - HIER STAND EINE ZEILE OHNE WIRKUNG:
                #   d.add_paragraph(f"— {cap}").italic = True
                # 'italic' ist keine Eigenschaft eines Absatzes, sondern
                # eines Laufs. python-docx legt auf dem Absatzobjekt
                # stillschweigend ein neues Attribut an, das niemand liest;
                # die Quellenangabe war im Word-Dokument NIE kursiv.
                # GEMESSEN am 13.08.2026 mit python-docx 1.2.0:
                #   p = d.add_paragraph("— Quelle"); p.italic = True
                #   -> getattr(p, "italic") == True, aber [r.italic for r
                #      in p.runs] == [None]
                # Das ist dieselbe Art Fehler, um die es in diesem Vorgang
                # geht: eine Zusage, die niemand einloest. Deshalb hier
                # mitbehoben - die Kursivstellung gehoert an den Lauf.
                quelle = d.add_paragraph()
                quelle.add_run(f"— {cap}").italic = True
        elif bt == "image":
            url = blk.data.get("_image_url", "")
            avail = blk.data.get("_image_available", False)
            p = d.add_paragraph()
            p.add_run("Bildverweis (nicht eingebettet — §§184b/184c): ").bold = True
            p.add_run(f"Quelle: {url or '(keine URL)'}; "
                      f"Status: {'vorhanden' if avail else 'NICHT auffindbar'}")
            if avail:
                uh = blk.data.get("_image_url_hash")
                aid = blk.data.get("_image_asset_id")
                size = blk.data.get("_image_size")
                bits = []
                if uh:
                    bits.append(f"url_hash={uh}")
                if aid is not None:
                    bits.append(f"asset_id={aid}")
                if size is not None:
                    bits.append(f"file_size={size} B")
                if bits:
                    d.add_paragraph("Anker: " + " · ".join(bits))
            cap = blk.data.get("_resolved_caption_plain", "")
            if cap:
                d.add_paragraph(f"Bildunterschrift: {cap}")
        elif bt == "delimiter":
            d.add_paragraph("———")
        elif bt == "marker":
            r = d.add_paragraph().add_run(blk.resolved_text_plain)
            r.bold = True
        elif bt == "evidence":
            if blk.resolved_text_plain:
                d.add_paragraph(blk.resolved_text_plain)
            ev_ids = blk.data.get("evidence_ids", [])
            if isinstance(ev_ids, list) and ev_ids:
                d.add_paragraph("Beweis-IDs: " + ", ".join(str(e) for e in ev_ids))

        self._anchors(d, blk, Pt)

    def _anchors(self, d, blk: RenderedBlock, Pt) -> None:
        if not blk.anchors:
            return
        p = d.add_paragraph()
        run = p.add_run("Beweisanker: ")
        run.bold = True
        run.font.size = Pt(9)
        for a in blk.anchors:
            aid = getattr(a, "id", "")
            atext = getattr(a, "anchor_text", "")
            p.add_run(f"[{aid}] {atext}  ").font.size = Pt(9)
